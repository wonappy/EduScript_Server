# utils/docx_generator.py
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from src.app.modules.file.file_data import FileData

# 강의 모드
def create_docx_from_text(text: str, filename: str, user_filename: str) -> FileData:
    """텍스트를 DOCX로 변환하여 FileData 객체 생성"""
    try:
        print(f"[DOCX DEBUG] DOCX 생성 시작 - 파일명: {filename}")
        print(f"[DOCX DEBUG] 사용자 파일명: {user_filename}")
        print(f"[DOCX DEBUG] 텍스트 길이: {len(text)} 문자")
        
        # Document 객체 생성
        doc = Document()
        
        # 문서 여백 설정 (선택사항)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 파일명에서 문서 타입 추출
        document_type = ""
        if "refined" in filename:
            document_type = " - 정제된 내용"
        elif "summary" in filename:
            document_type = " - 요약"
        elif "key_points" in filename:
            document_type = " - 핵심 포인트"
        
        # 제목 생성 및 추가
        title = user_filename + document_type
        print(f"[DOCX DEBUG] 제목: '{title}'")
        
        # 제목 스타일 추가
        title_paragraph = doc.add_heading(title, 0)  # 0 = 가장 큰 제목
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 빈 줄 추가
        doc.add_paragraph()
        
        print(f"[DOCX DEBUG] 텍스트 처리 시작")
        
        # 본문 처리
        paragraphs = text.split('\n')
        total_paragraphs = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            if not paragraph.strip():
                # 빈 줄 처리
                doc.add_paragraph()
                continue
            
            # 문단 추가
            p = doc.add_paragraph()
            run = p.add_run(paragraph)
            
            # 폰트 설정 (선택사항)
            run.font.size = Pt(12)  # 12pt 폰트
            
            total_paragraphs += 1
            
            if para_idx < 5:  # 처음 5개 문단만 로그 출력
                print(f"[DOCX DEBUG] 문단 {para_idx+1}: '{paragraph[:30]}...'")
        
        print(f"[DOCX DEBUG] 텍스트 처리 완료 - 총 {total_paragraphs}개 문단")
        
        # BytesIO에 저장
        buffer = BytesIO()
        doc.save(buffer)
        
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"[DOCX DEBUG] DOCX 크기: {len(docx_bytes)} bytes")
        
        # DOCX 헤더 확인 (ZIP 파일 형태)
        if not docx_bytes.startswith(b'PK'):
            raise Exception("생성된 DOCX에 올바른 헤더가 없습니다")
        
        print(f"[DOCX DEBUG] DOCX 헤더 검증 완료")
        
        # FileData 객체 생성
        file_data = FileData.from_bytes(
            data=docx_bytes,
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        print(f"[DOCX DEBUG] FileData 생성 완료 - 파일크기: {file_data.file_size}")
        return file_data
        
    except Exception as e:
        print(f"[DOCX ERROR] DOCX 생성 실패: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"DOCX 생성 실패: {str(e)}")

# 회의 모드
def create_meeting_minutes_docx(json_data: str, filename: str, user_filename: str) -> FileData:
    """JSON 데이터를 받아 회의록 양식 DOCX로 변환"""
    try:
        print(f"[회의록 DOCX] 생성 시작 - 파일명: {filename}")
        
        # JSON 파싱
        try:
            if isinstance(json_data, str):
                data = json.loads(json_data)
            elif isinstance(json_data, dict):
                data = json_data
            else:
                data = json.loads(str(json_data))
                
            print(f"[회의록 DOCX] JSON 파싱 완료")
            
        except Exception as parse_error:
            print(f"[회의록 DOCX] JSON 파싱 실패: {parse_error}")
            # 기본 빈 구조
            data = {
                "회의안건": [],
                "회의내용": [],
                "결정사항": [],
                "특이사항": ""
            }
        
        # Document 객체 생성
        doc = Document()
        
        # 페이지 설정 (간소화)
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # 제목
        title = f"{user_filename} - 회의록"
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 빈 줄
        doc.add_paragraph()
        
        # === 기본 정보 테이블 ===
        info_table = doc.add_table(rows=2, cols=3)
        info_table.style = 'Light Grid Accent 1'  # 안전한 기본 스타일
        
        # 헤더
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = '회의일시'
        hdr_cells[1].text = '부서'
        hdr_cells[2].text = '작성자'
        
        # 데이터 행
        data_cells = info_table.rows[1].cells
        data_cells[0].text = '20    년    월    일'
        data_cells[1].text = ''
        data_cells[2].text = ''
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # === 참석자 ===
        doc.add_heading('참석자', level=2)
        doc.add_paragraph()  # 빈 공간
        doc.add_paragraph()
        doc.add_paragraph()
        
        # === 회의안건 ===
        doc.add_heading('회의안건', level=2)
        
        agenda_items = data.get('회의안건', [])
        num_agenda_rows = max(3, len(agenda_items)) if agenda_items else 3
        
        agenda_table = doc.add_table(rows=num_agenda_rows + 1, cols=1)
        agenda_table.style = 'Table Grid'
        
        # 헤더
        agenda_table.rows[0].cells[0].text = '안건'
        
        # 내용
        for i in range(num_agenda_rows):
            row_idx = i + 1
            if i < len(agenda_items) and agenda_items[i]:
                agenda_table.rows[row_idx].cells[0].text = f"{i+1}. {str(agenda_items[i])}"
            else:
                agenda_table.rows[row_idx].cells[0].text = f"{i+1}. "
        
        doc.add_paragraph()
        
        # === 회의내용 ===
        doc.add_heading('회의내용', level=2)
        
        content_items = data.get('회의내용', [])
        num_content_rows = max(3, len(content_items)) if content_items else 3
        
        content_table = doc.add_table(rows=num_content_rows + 1, cols=2)
        content_table.style = 'Table Grid'
        
        # 헤더
        content_table.rows[0].cells[0].text = '내용'
        content_table.rows[0].cells[1].text = '비고'
        
        # 내용
        for i in range(num_content_rows):
            row_idx = i + 1
            if i < len(content_items) and content_items[i]:
                item = content_items[i]
                if isinstance(item, dict):
                    content_table.rows[row_idx].cells[0].text = str(item.get('내용', ''))
                    content_table.rows[row_idx].cells[1].text = str(item.get('비고', ''))
                else:
                    content_table.rows[row_idx].cells[0].text = str(item)
                    content_table.rows[row_idx].cells[1].text = ''
            else:
                content_table.rows[row_idx].cells[0].text = ''
                content_table.rows[row_idx].cells[1].text = ''
        
        doc.add_paragraph()
        
        # === 결정사항 ===
        doc.add_heading('결정사항', level=2)
        
        decision_items = data.get('결정사항', [])
        num_decision_rows = max(3, len(decision_items)) if decision_items else 3
        
        decision_table = doc.add_table(rows=num_decision_rows + 1, cols=2)
        decision_table.style = 'Table Grid'
        
        # 헤더
        decision_table.rows[0].cells[0].text = '내용'
        decision_table.rows[0].cells[1].text = '진행일정'
        
        # 내용
        for i in range(num_decision_rows):
            row_idx = i + 1
            if i < len(decision_items) and decision_items[i]:
                item = decision_items[i]
                if isinstance(item, dict):
                    decision_table.rows[row_idx].cells[0].text = str(item.get('내용', ''))
                    decision_table.rows[row_idx].cells[1].text = str(item.get('진행일정', ''))
                else:
                    decision_table.rows[row_idx].cells[0].text = str(item)
                    decision_table.rows[row_idx].cells[1].text = ''
            else:
                decision_table.rows[row_idx].cells[0].text = ''
                decision_table.rows[row_idx].cells[1].text = ''
        
        doc.add_paragraph()
        
        # === 특이사항 ===
        doc.add_heading('특이사항', level=2)
        
        special_content = data.get('특이사항', '')
        if special_content and str(special_content).strip():
            doc.add_paragraph(str(special_content))
        else:
            # 빈 공간
            for _ in range(4):
                doc.add_paragraph()
        
        print(f"[회의록 DOCX] 문서 구조 생성 완료")
        
        # BytesIO에 저장
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # 포인터를 처음으로 이동
        
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"[회의록 DOCX] 바이트 생성 완료 - 크기: {len(docx_bytes)} bytes")
        
        # DOCX 파일 헤더 확인
        if len(docx_bytes) == 0:
            raise Exception("생성된 DOCX 파일이 비어있습니다")
            
        if not docx_bytes.startswith(b'PK'):
            raise Exception(f"올바르지 않은 DOCX 헤더: {docx_bytes[:10]}")
        
        print(f"[회의록 DOCX] 헤더 검증 통과")
        
        # FileData 객체 생성
        file_data = FileData.from_bytes(
            data=docx_bytes,
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        print(f"[회의록 DOCX] FileData 생성 완료 - 파일크기: {file_data.file_size}")
        return file_data
        
    except Exception as e:
        print(f"[회의록 DOCX ERROR] 생성 실패: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"회의록 DOCX 생성 실패: {str(e)}")